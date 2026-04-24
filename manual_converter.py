import sys
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar

# ==========================================
# 📝 QUICK EDIT: ADD YOUR FILE NAME HERE 
# ==========================================
MY_FILE = "Sales Engineer Job Profile.docx"  # <--- Change this to your file name!

# ==========================================
# SETTINGS (EXTRACTED FROM INTERNSHIP MANUAL)
# ==========================================
DEFAULT_FONT = "Cambria"
BODY_SIZE = 12
HEADER_SIZE = 14
MARGINS = 1.0
THEME_COLOR = "4D4D80" # The dark blue used in table headers

def add_page_border(section):
    """Adds a black rectangular border to a section."""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for b_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{b_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 1/2 pt
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)

def set_cell_bg(cell, hex_color):
    """Sets the background color of a table cell."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def format_word_file(input_path):
    """Applies all manual formatting to a docx file."""
    if not input_path.endswith('.docx'):
        print(f"Skipping {input_path}: Not a .docx file.")
        return

    print(f"Processing: {input_path}...")
    doc = Document(input_path)
    
    # 1. Margins and Borders
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(MARGINS)
        section.left_margin = section.right_margin = Inches(MARGINS)
        add_page_border(section)

    # 2. Text Formatting
    for para in doc.paragraphs:
        # Simple heuristic for headers (bold text)
        is_bold = any(run.bold for run in para.runs)
        for run in para.runs:
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(HEADER_SIZE if is_bold else BODY_SIZE)
        para.paragraph_format.line_spacing = 1.15

    # 3. Table Formatting
    for table in doc.tables:
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                set_cell_bg(cell, THEME_COLOR)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.bold = True
                        r.font.name = DEFAULT_FONT
                        r.font.size = Pt(11)

    output_path = "formatted_" + os.path.basename(input_path)
    doc.save(output_path)
    print(f"✅ Success! Created: {output_path}")

def analyze_pdf_styles(pdf_path):
    """Extra: Analyze a PDF if you want to learn new styles."""
    print(f"Analyzing PDF: {pdf_path}...")
    # ... logic from previous analyze scripts ...
    # This is here just so you have all code in one file
    pass

def merge_and_format_files(input_paths, output_path):
    """Combines multiple Word files into one and applies formatting."""
    print(f"Merging {len(input_paths)} files into {output_path}...")
    
    master_doc = Document()
    
    for i, path in enumerate(input_paths):
        if not path.endswith('.docx'):
            continue
        
        sub_doc = Document(path)
        
        # Add page break between files except for the last one
        if i > 0:
            master_doc.add_page_break()
            
        # Copy content placeholder (simplified: docxcompose would be better, 
        # but we can copy paragraphs/tables manually for a basic version)
        for element in sub_doc.element.body:
            master_doc.element.body.append(element)

    # Save temporary merged document
    master_doc.save("temp_merged.docx")
    
    # Apply global formatting to the merged result
    format_word_file("temp_merged.docx")
    
    # Rename result
    if os.path.exists("formatted_temp_merged.docx"):
        if os.path.exists(output_path):
            os.remove(output_path)
        os.rename("formatted_temp_merged.docx", output_path)
        os.remove("temp_merged.docx")
        print(f"✅ Final merged and formatted file created: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        # Check if the user specified a file in the QUICK EDIT section
        if MY_FILE and MY_FILE != "change_me.docx":
            if os.path.exists(MY_FILE):
                format_word_file(MY_FILE)
            else:
                print(f"Error: File '{MY_FILE}' was not found in this folder.")
                print("Make sure you put your file in: " + os.getcwd())
        else:
            # AUTOMATIC MODE: Find all .docx files in the current directory
            print("--- Manual Converter Auto-Mode ---")
            docx_files = [f for f in os.listdir('.') if f.endswith('.docx') and not f.startswith('formatted_')]
            
            if not docx_files:
                print("No .docx files found in this folder!")
            else:
                for f in docx_files:
                    format_word_file(f)
            
    elif sys.argv[1] == "--merge":
        if len(sys.argv) < 4:
            print("Error: Provide output name and at least one input file.")
        else:
            out_name = sys.argv[2]
            in_files = sys.argv[3:]
            merge_and_format_files(in_files, out_name)
    else:
        for file_arg in sys.argv[1:]:
            if os.path.exists(file_arg):
                format_word_file(file_arg)
            else:
                print(f"Error: File '{file_arg}' not found.")
